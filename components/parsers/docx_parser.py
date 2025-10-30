# /Users/murseltasgin/projects/chat_rag/components/parsers/docx_parser.py
"""
Microsoft Word document parser
"""
import os
from typing import Dict, Any
from .base import BaseParser
from core.exceptions import RAGException


class DOCXParser(BaseParser):
    """Parser for Microsoft Word documents (.docx)"""
    
    def __init__(self):
        """Initialize DOCX parser"""
        # Try to import required library
        try:
            import docx
            self.docx = docx
        except ImportError:
            raise RAGException(
                "python-docx is required for DOCX parsing. Install with: pip install python-docx"
            )
    
    def parse(self, file_path: str, **kwargs) -> str:
        """Parse DOCX file"""
        try:
            doc = self.docx.Document(file_path)
            
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise RAGException(f"Failed to parse DOCX file {file_path}: {e}")
    
    def supports(self, file_path: str) -> bool:
        """Check if file is a DOCX"""
        ext = os.path.splitext(file_path)[1].lower()
        return ext == '.docx'
    
    def get_name(self) -> str:
        """Get parser name"""
        return "DOCXParser"
    
    def get_metadata(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Extract metadata from DOCX"""
        metadata = super().get_metadata(file_path)
        
        try:
            doc = self.docx.Document(file_path)
            
            # Count elements
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            
            # Extract core properties
            if doc.core_properties.title:
                metadata['title'] = doc.core_properties.title
            if doc.core_properties.author:
                metadata['author'] = doc.core_properties.author
            if doc.core_properties.subject:
                metadata['subject'] = doc.core_properties.subject
            if doc.core_properties.created:
                metadata['created'] = str(doc.core_properties.created)
            if doc.core_properties.modified:
                metadata['modified'] = str(doc.core_properties.modified)
        except Exception:
            pass
        
        return metadata

